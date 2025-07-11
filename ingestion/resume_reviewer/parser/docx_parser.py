from docx import Document


def _iter_table_text(doc: Document):
    """Yield non-empty cell text from every table in the document."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    yield txt


def _iter_header_footer_text(doc: Document):
    """Yield text from all section headers and footers."""
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                txt = p.text.strip()
                if txt:
                    yield txt


def extract(path: str) -> str:
    doc = Document(path)

    lines: list[str] = []

    # --- body paragraphs ---------------------------------------------------
    lines.extend(p.text.strip() for p in doc.paragraphs if p.text.strip())

    # --- tables -------------------------------------------------------------
    lines.extend(_iter_table_text(doc))

    # --- headers & footers --------------------------------------------------
    lines.extend(_iter_header_footer_text(doc))

    return "\n".join(lines)
